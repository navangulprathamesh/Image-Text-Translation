import pytesseract
import cv2
import numpy as np
from pytesseract import Output
import boto3
import img2pdf
import os
import sys
import re
import io
import glob
from PIL import Image
from PIL import ImageFont
from PIL import ImageDraw
from docx import Document
from pdf2image import convert_from_path
from PIL import Image
import img2pdf
from google_trans_new import google_translator  




image_path = "Sir_Image.pdf"
L1=list(os.path.splitext(image_path))
image1=np.zeros((800,1000,3), np.uint8)
if L1[1]== '.pdf':
    images = convert_from_path(image_path, 500,poppler_path=r'C:\Users\Abhi\Downloads\poppler-0.68.0\bin')
    for i in range(len(images)):
        print("Page "+str(i+1)+" processing.....")
        
        
        images[i].save('pdf/output.jpg')
        image=cv2.imread("pdf/output.jpg")
        image=cv2.resize(image,(1000,800))
        with open("pdf/output.jpg", 'rb') as document:
            imageBytes = bytearray(document.read())
        text = boto3.client('textract')
        response = text.detect_document_text(Document={'Bytes': imageBytes})
        locations=[]
        words=[]
        for j in response["Blocks"]:
            if j["BlockType"] == "LINE":
                words.append('{}'.format(j["Text"]))
                locations.append([j["Geometry"]["BoundingBox"]["Width"],j["Geometry"]["BoundingBox"]["Height"],j["Geometry"]["BoundingBox"]["Left"],j["Geometry"]["BoundingBox"]["Top"]])

        text="  ".join(words)
        
        text=list(map(str,text.split("  ")))
        trans=[]
        '''
        translate = boto3.client(service_name='translate', region_name='us-east-2', use_ssl=True)
        for j in text:
            result = translate.translate_text(Text=j, SourceLanguageCode="en", TargetLanguageCode="hi")
            trans.append(result.get("TranslatedText"))


        #translated text
        translator=" ".join(trans)
        translator=list(map(str,translator.split("  ")))
        '''
        translator = google_translator() 
        for j in text:
            trans.append(translator.translate(j,lang_tgt='mr'))
       
        print("Translated Text: ")
        print(" ".join(trans))
        
        list2=[]
        for j in locations:
            (x, y, w, h) = (j[2], j[3], j[0],  j[1])
            image = cv2.rectangle(image, (int(x*1000), int(y*800)), (int((x+w)*1000), int((y+h)*800)), (255, 255, 255), -1)
            image1 = cv2.rectangle(image1, (int(x*1000), int(y*800)), (int((x+w)*1000), int((y+h)*800)), (255, 255, 255), -1)
            list2.append([int(x*1000), int(y*800)])
        
      
        cv2.imwrite("pdf/text1.jpg",image)
        cv2.imwrite("pdf/text2.jpg",image1)
        image=cv2.imread("pdf/text1.jpg")
        image1=cv2.imread("pdf/text2.jpg",0)
        dst = cv2.inpaint(image, image1, 2, cv2.INPAINT_TELEA)
        cv2.imwrite("pdf/text3.jpg",dst)
        final = Image.open("pdf/text3.jpg")


        #Language to be translated
        if lang_tgt=='mr':
            font = ImageFont.truetype("Marathi.ttf", 18)
        elif lang_tgt=='pa' :
            font = ImageFont.truetype("Karmic_Sanj_Black.ttf", 16)
        elif lang_tgt=='hi' :
            font = ImageFont.truetype("Akshar Unicode.ttf", 19)

        draw = ImageDraw.Draw(final)

        for j in range(len(trans)):
            draw.text( (list2[j][0],list2[j][1]), trans[j], fill=(0,0,0),font=font)
        final.save("pdf/Final"+str(i)+".jpg")
        os.remove("pdf/text1.jpg")
        os.remove("pdf/text2.jpg")
        os.remove("pdf/text3.jpg")
        os.remove("pdf/output.jpg")
        print("Image "+str(i+1)+" Inpainted successfully.\n")
    with open("final/Final.pdf","wb") as f:
        f.write(img2pdf.convert(glob.glob("pdf/*.jpg")))
    for file in os.scandir("pdf/"):
        os.remove(file.path)
    print("Final PDF Generated Successfully.")
        

    
elif L1[1]== '.png' or L1[1]== '.jpg' or L1[1]== '.jpeg':
    image = cv2.imread(image_path)
    image=cv2.resize(image,(1000,800))
    cv2.imwrite("self.png",image)
    
    with open(image_path, 'rb') as document:
        imageBytes = bytearray(document.read())
    text = boto3.client('textract')
    response = text.detect_document_text(Document={'Bytes': imageBytes})
    locations=[]
    words=[]
    for i in response["Blocks"]:
        if i["BlockType"] == "LINE":
            words.append('{}'.format(i["Text"]))
            locations.append([i["Geometry"]["BoundingBox"]["Width"],i["Geometry"]["BoundingBox"]["Height"],i["Geometry"]["BoundingBox"]["Left"],i["Geometry"]["BoundingBox"]["Top"]])
    #extracted text
    text="  ".join(words)
    print(" ".join(words))

    text=list(map(str,text.split("  ")))
    trans=[]
    '''
    translate = boto3.client(service_name='translate', region_name='us-east-2', use_ssl=True)
    for i in text:
        result = translate.translate_text(Text=i, SourceLanguageCode="en", TargetLanguageCode="hi")
        trans.append(result.get("TranslatedText"))
    print(" ".join(trans))

    #translated text
    translator=" ".join(trans)
    translator=list(map(str,translator.split("  ")))
    '''

    translator = google_translator() 
        for j in text:
            trans.append(translator.translate(j,lang_tgt='mr'))
       
        print("Translated Text: ")
        print(" ".join(trans))
    
    list2=[]
    for i in locations:
        (x, y, w, h) = (i[2], i[3], i[0],  i[1])
        image = cv2.rectangle(image, (int(x*1000), int(y*800)), (int((x+w)*1000), int((y+h)*800)), (255, 255, 255), -1)
        image1 = cv2.rectangle(image1, (int(x*1000), int(y*800)), (int((x+w)*1000), int((y+h)*800)), (255, 255, 255), -1)
        list2.append([int(x*1000), int(y*800)])

    #print(list2)
    #cv2.imshow('Img',image1) 
    cv2.imwrite("text1.jpg",image)
    cv2.imwrite("text2.jpg",image1)
    image=cv2.imread("text1.jpg")
    image1=cv2.imread("text2.jpg",0)
    dst = cv2.inpaint(image, image1, 2, cv2.INPAINT_TELEA)
    cv2.imwrite("text3.jpg",dst)
    final = Image.open("text3.jpg")

    #Language to be translated
    if lang_tgt='mr' :
        font = ImageFont.truetype("Marathi.ttf", 18)
    elif lang_tgt='pa' :
        font = ImageFont.truetype("Karmic_Sanj_Black.ttf", 16)
    elif lang_tgt='hi' :
        font = ImageFont.truetype("Akshar Unicode.ttf", 19)

    draw = ImageDraw.Draw(final)
    cv2.imshow('Inpainting result', dst)

    for i in range(len(trans)):
        draw.text( (list2[i][0],list2[i][1]), trans[i], fill=(0,0,0),font=font)
    final.save("Final2.jpg")



    """img_path = "Final2.jpg"
    pdf_path = "file.pdf"
    image = Image.open(img_path)
    pdf_bytes = img2pdf.convert(image.filename)
    file = open(pdf_path, "wb")
    file.write(pdf_bytes)
    image.close()
    file.close()"""
    #print("Successfully made pdf file")

    #path='file1.txt'
    #F = open(path,"w")
    """with io.open(path,"w", encoding="utf-8") as F:
        F.write('Extracted:\n')
        F.write("\n".join([str(x) for x in text]))
        F.write('\n\n')
        F.write('Translated:\n')
        F.write("\n".join([str(x) for x in translator]))"""
        # F.seek(0)
        #F.close()
   # print("Successfully made text file")
    
    """path1='E:/Btech cse/Btech OCR project/mega-project/Hand written recog/Word_file/file1'
    print(path1)
    direct = os.listdir(path1)
    for i in direct:
        document = Document()
        document.add_heading(i, 0)
        myfile = open('/path/to/read/from/'+i).read()
        myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
        p = document.add_paragraph(myfile)
        document.save('/path/to/write/to/'+i+'.docx')"""
    print("Done")
    cv2.waitKey(0)
    cv2.destroyAllWindows()






